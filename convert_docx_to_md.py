#!/usr/bin/env python3
"""
Convert DOCX file to Markdown format
"""

from docx import Document
import re
import os

def convert_docx_to_markdown(docx_path, md_path):
    """Convert a DOCX file to Markdown format"""
    
    try:
        # Load the Word document
        doc = Document(docx_path)
        
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                markdown_content.append("")
                continue
            
            # Check for heading styles
            style_name = paragraph.style.name.lower()
            
            if 'heading 1' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}")
            else:
                # Regular paragraph
                # Check for bold and italic formatting in runs
                formatted_text = ""
                for run in paragraph.runs:
                    run_text = run.text
                    if run.bold and run.italic:
                        formatted_text += f"***{run_text}***"
                    elif run.bold:
                        formatted_text += f"**{run_text}**"
                    elif run.italic:
                        formatted_text += f"*{run_text}*"
                    else:
                        formatted_text += run_text
                
                # If no formatting was applied, use the original text
                if not any(run.bold or run.italic for run in paragraph.runs):
                    formatted_text = text
                
                markdown_content.append(formatted_text)
        
        # Handle tables
        for table in doc.tables:
            markdown_content.append("")  # Add empty line before table
            
            # Process header row
            if table.rows:
                header_row = table.rows[0]
                header_cells = [cell.text.strip() for cell in header_row.cells]
                
                # Create header
                markdown_content.append("| " + " | ".join(header_cells) + " |")
                markdown_content.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                
                # Process data rows
                for row in table.rows[1:]:
                    data_cells = [cell.text.strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(data_cells) + " |")
            
            markdown_content.append("")  # Add empty line after table
        
        # Write to markdown file
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        print(f"✅ Successfully converted '{docx_path}' to '{md_path}'")
        print(f"📄 Content preview:")
        print("-" * 50)
        
        # Show first few lines as preview
        preview_lines = markdown_content[:15]
        for line in preview_lines:
            print(line)
        
        if len(markdown_content) > 15:
            print(f"\n... and {len(markdown_content) - 15} more lines")
        
        return True
        
    except Exception as e:
        print(f"❌ Error converting file: {e}")
        return False

if __name__ == "__main__":
    # File paths
    docx_file = "last_decisions.docx"
    md_file = "last_decisions.md"
    
    if os.path.exists(docx_file):
        print(f"🔄 Converting {docx_file} to {md_file}...")
        convert_docx_to_markdown(docx_file, md_file)
    else:
        print(f"❌ File {docx_file} not found!")
